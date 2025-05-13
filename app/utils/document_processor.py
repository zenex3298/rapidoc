import os
import re
import shutil
import json
import logging
import numpy as np
import time
import uuid
from pathlib import Path
from typing import Dict, Any, List, Optional, Union
from datetime import datetime
from dotenv import load_dotenv

# Try to import various document processing libraries
# If any fail, we'll handle it gracefully
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

try:
    import boto3
    BOTO3_AVAILABLE = True
except ImportError:
    BOTO3_AVAILABLE = False

# Set up logging
logger = logging.getLogger(__name__)

# Load environment variables for S3 configuration
load_dotenv()

# Configure S3 if enabled via environment variable
use_s3 = os.getenv('USE_S3_STORAGE', 'false').lower() == 'true'
s3_client = None
s3_bucket = os.getenv('S3_BUCKET_NAME') or os.getenv('S3_BUCKET')

if use_s3:
    if not BOTO3_AVAILABLE:
        logger.warning("S3 storage enabled but boto3 is not installed. Install with: pip install boto3")
    else:
        try:
            s3_client = boto3.client('s3',
                aws_access_key_id=os.getenv('AWS_ACCESS_KEY_ID'),
                aws_secret_access_key=os.getenv('AWS_SECRET_ACCESS_KEY'),
                region_name=os.getenv('AWS_REGION', 'us-east-1')
            )
            if s3_bucket:
                logger.info(f"S3 client initialized for bucket: {s3_bucket}")
            else:
                logger.warning("S3 client initialized but no bucket name found in S3_BUCKET_NAME or S3_BUCKET environment variables")
        except Exception as e:
            logger.error(f"Failed to initialize S3 client: {e}")
            s3_client = None


def convert_numpy_to_python(obj):
    """Convert NumPy types to Python native types for JSON serialization.
    
    Args:
        obj: Object potentially containing NumPy types
        
    Returns:
        Object with NumPy types converted to Python native types
    """
    if isinstance(obj, np.integer):
        return int(obj)
    elif isinstance(obj, np.floating):
        return float(obj)
    elif isinstance(obj, np.ndarray):
        return obj.tolist()
    elif isinstance(obj, dict):
        return {k: convert_numpy_to_python(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [convert_numpy_to_python(i) for i in obj]
    else:
        return obj


class DocumentProcessor:
    """Processor for document extraction and analysis."""

    def __init__(self, upload_dir: str = "uploads"):
        """Initialize the document processor.
        
        Args:
            upload_dir: Directory to store uploaded files
        """
        self.upload_dir = upload_dir
        
        # Create upload directory if it doesn't exist (for local storage)
        if not use_s3:
            os.makedirs(upload_dir, exist_ok=True)
            logger.info(f"Using local storage at: {upload_dir}")
        else:
            logger.info(f"Using S3 storage with bucket: {s3_bucket}")

    def save_file(self, file_content: bytes, filename: str, user_id: int) -> str:
        """Save a file to the appropriate storage (local or S3).
        
        Args:
            file_content: The file content as bytes
            filename: The original filename
            user_id: ID of the user uploading the file
            
        Returns:
            str: Path where the file is saved (local path or S3 key)
        """
        # Create a safe filename
        current_time = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_filename = re.sub(r'[^\w\.-]', '_', filename)
        safe_filename = f"{current_time}_{safe_filename}"
        
        # Create user-specific directory
        user_dir = os.path.join(self.upload_dir, str(user_id))
        
        if use_s3 and s3_client:
            # Save to S3
            try:
                # Check if bucket name is configured
                if not s3_bucket:
                    raise ValueError("S3 bucket name not configured. Check S3_BUCKET_NAME environment variable.")
                    
                object_key = f"{user_dir}/{safe_filename}"
                s3_client.put_object(
                    Bucket=s3_bucket,
                    Key=object_key,
                    Body=file_content
                )
                logger.info(f"Saved file to S3: {object_key}")
                return object_key
            except Exception as e:
                logger.error(f"Error saving file to S3: {e}")
                raise
        else:
            # Save locally
            try:
                # Create user directory if it doesn't exist
                os.makedirs(user_dir, exist_ok=True)
                
                # Full path to save the file
                file_path = os.path.join(user_dir, safe_filename)
                
                # Write the file
                with open(file_path, 'wb') as f:
                    f.write(file_content)
                    
                logger.info(f"Saved file locally: {file_path}")
                return file_path
            except Exception as e:
                logger.error(f"Error saving file locally: {e}")
                raise

    def save_transformed_file(self, file_content: str, file_type: str, original_document_title: str, user_id: int) -> str:
        """Save a transformed file (text content) to the appropriate storage.
        
        Args:
            file_content: The file content as string
            file_type: The file type/extension (without the dot)
            original_document_title: Title of the original document
            user_id: ID of the user performing the transformation
            
        Returns:
            str: Path where the file is saved (local path or S3 key)
        """
        # Create a filename based on the original document title
        current_time = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_title = re.sub(r'[^\w\.-]', '_', original_document_title)
        safe_filename = f"{current_time}_{safe_title}_transformed.{file_type}"
        
        # Create user-specific directory
        user_dir = os.path.join(self.upload_dir, str(user_id))
        
        if use_s3 and s3_client:
            # Save to S3
            try:
                # Check if bucket name is configured
                if not s3_bucket:
                    raise ValueError("S3 bucket name not configured. Check S3_BUCKET_NAME environment variable.")
                    
                object_key = f"{user_dir}/{safe_filename}"
                s3_client.put_object(
                    Bucket=s3_bucket,
                    Key=object_key,
                    Body=file_content.encode('utf-8')  # Convert string to bytes
                )
                logger.info(f"Saved transformed file to S3: {object_key}")
                return object_key
            except Exception as e:
                logger.error(f"Error saving transformed file to S3: {e}")
                raise
        else:
            # Save locally
            try:
                # Create user directory if it doesn't exist
                os.makedirs(user_dir, exist_ok=True)
                
                # Full path to save the file
                file_path = os.path.join(user_dir, safe_filename)
                
                # Write the file
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(file_content)
                    
                logger.info(f"Saved transformed file locally: {file_path}")
                return file_path
            except Exception as e:
                logger.error(f"Error saving transformed file locally: {e}")
                raise

    def get_file_content(self, file_path: str) -> bytes:
        """Get file content from the appropriate storage.
        
        Args:
            file_path: Path to the file (local path or S3 key)
            
        Returns:
            bytes: The file content
        """
        if use_s3 and s3_client and not os.path.exists(file_path):
            # File is in S3
            try:
                # Check if bucket name is configured
                if not s3_bucket:
                    raise ValueError("S3 bucket name not configured. Check S3_BUCKET_NAME environment variable.")
                    
                response = s3_client.get_object(
                    Bucket=s3_bucket,
                    Key=file_path
                )
                content = response['Body'].read()
                logger.info(f"Retrieved file from S3: {file_path}")
                return content
            except Exception as e:
                logger.error(f"Error retrieving file from S3: {e}")
                raise
        else:
            # File is local
            try:
                with open(file_path, 'rb') as f:
                    content = f.read()
                logger.info(f"Retrieved file locally: {file_path}")
                return content
            except Exception as e:
                logger.error(f"Error retrieving file locally: {e}")
                raise

    def extract_text(self, file_path: str) -> str:
        """Extract text content from a document file.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            str: Extracted text content
        """
        try:
            # Start timing for performance logging
            start_time = time.time()
            
            # Get file extension
            if '.' in os.path.basename(file_path):
                file_ext = os.path.splitext(file_path)[1].lower()
            else:
                # Handle files without extensions
                # Try to infer type from first few bytes
                file_content = self.get_file_content(file_path)
                file_ext = self._infer_file_type(file_content)
                
            # Process based on file type
            if file_ext == '.pdf':
                text = self._extract_text_from_pdf(file_path)
            elif file_ext in ['.doc', '.docx']:
                text = self._extract_text_from_word(file_path)
            elif file_ext in ['.csv']:
                text = self._extract_text_from_csv(file_path)
            elif file_ext in ['.xls', '.xlsx']:
                text = self._extract_text_from_excel(file_path)
            elif file_ext in ['.txt', '.md', '.json', '.html', '.htm', '.xml']:
                # Plain text files
                file_content = self.get_file_content(file_path)
                try:
                    text = file_content.decode('utf-8')
                except UnicodeDecodeError:
                    # Try another encoding if UTF-8 fails
                    text = file_content.decode('latin-1')
            else:
                # Unsupported format
                logger.warning(f"Unsupported file format: {file_ext}")
                text = f"Unsupported file format: {file_ext}"
            
            # Log performance metrics
            processing_time = time.time() - start_time
            logger.info(f"Text extraction completed in {processing_time:.2f} seconds for file: {file_path}")
            
            return text
            
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return f"Error extracting text: {str(e)}"

    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from a document file.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dict[str, Any]: Extracted metadata
        """
        try:
            # Get file extension
            if '.' in os.path.basename(file_path):
                file_ext = os.path.splitext(file_path)[1].lower()
            else:
                # Handle files without extensions
                # Try to infer type from first few bytes
                file_content = self.get_file_content(file_path)
                file_ext = self._infer_file_type(file_content)
            
            # Basic metadata
            metadata = {
                "file_name": os.path.basename(file_path),
                "file_extension": file_ext,
                "file_size_bytes": self._get_file_size(file_path),
                "extraction_time": datetime.now().isoformat(),
            }
            
            # Add specific metadata by file type
            if file_ext == '.pdf' and PYPDF2_AVAILABLE:
                pdf_metadata = self._extract_pdf_metadata(file_path)
                metadata.update(pdf_metadata)
            elif file_ext in ['.doc', '.docx'] and DOCX_AVAILABLE:
                word_metadata = self._extract_word_metadata(file_path)
                metadata.update(word_metadata)
            elif file_ext in ['.csv'] and PANDAS_AVAILABLE:
                csv_metadata = self._extract_csv_metadata(file_path)
                metadata.update(csv_metadata)
            elif file_ext in ['.xls', '.xlsx'] and PANDAS_AVAILABLE:
                excel_metadata = self._extract_excel_metadata(file_path)
                metadata.update(excel_metadata)
            
            return metadata
            
        except Exception as e:
            logger.error(f"Error extracting metadata from {file_path}: {e}")
            return {"error": str(e)}

    def get_document_structure(self, file_path: str) -> Dict[str, Any]:
        """Analyze the document's structure.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dict[str, Any]: Document structure information
        """
        try:
            # Get file extension
            if '.' in os.path.basename(file_path):
                file_ext = os.path.splitext(file_path)[1].lower()
            else:
                # Handle files without extensions
                file_content = self.get_file_content(file_path)
                file_ext = self._infer_file_type(file_content)
            
            structure = {
                "file_type": file_ext
            }
            
            if file_ext == '.pdf' and PYPDF2_AVAILABLE:
                pdf_structure = self._analyze_pdf_structure(file_path)
                structure.update(pdf_structure)
            elif file_ext in ['.doc', '.docx'] and DOCX_AVAILABLE:
                word_structure = self._analyze_word_structure(file_path)
                structure.update(word_structure)
            elif file_ext in ['.csv'] and PANDAS_AVAILABLE:
                csv_structure = self._analyze_csv_structure(file_path)
                structure.update(csv_structure)
            elif file_ext in ['.xls', '.xlsx'] and PANDAS_AVAILABLE:
                excel_structure = self._analyze_excel_structure(file_path)
                structure.update(excel_structure)
            
            return structure
            
        except Exception as e:
            logger.error(f"Error analyzing document structure for {file_path}: {e}")
            return {"error": str(e)}

    def _get_file_size(self, file_path: str) -> int:
        """Get the file size in bytes.
        
        Args:
            file_path: Path to the file
            
        Returns:
            int: File size in bytes
        """
        if use_s3 and s3_client and not os.path.exists(file_path):
            # Get file size from S3
            try:
                # Check if bucket name is configured
                if not s3_bucket:
                    raise ValueError("S3 bucket name not configured. Check S3_BUCKET_NAME environment variable.")
                    
                response = s3_client.head_object(
                    Bucket=s3_bucket,
                    Key=file_path
                )
                return response['ContentLength']
            except Exception as e:
                logger.error(f"Error getting file size from S3: {e}")
                return 0
        else:
            # Get file size locally
            try:
                return os.path.getsize(file_path)
            except Exception as e:
                logger.error(f"Error getting file size locally: {e}")
                return 0

    def _infer_file_type(self, file_content: bytes) -> str:
        """Try to infer file type from content.
        
        Args:
            file_content: File content as bytes
            
        Returns:
            str: Inferred file extension
        """
        # Check PDF signature
        if file_content.startswith(b'%PDF'):
            return '.pdf'
        
        # Check ZIP-based formats (DOCX, XLSX)
        if file_content.startswith(b'PK\x03\x04'):
            # This could be DOCX, XLSX, PPTX, etc.
            # Without deeper inspection, we'll default to DOCX
            return '.docx'
        
        # Check for CSV (look for commas and newlines)
        if b',' in file_content and b'\n' in file_content:
            return '.csv'
        
        # Default to txt for unknown types
        return '.txt'

    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from a PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            str: Extracted text
        """
        if not PYPDF2_AVAILABLE:
            return "Error: PyPDF2 library not available for PDF extraction"
        
        text_parts = []
        
        try:
            # Get file content
            file_content = self.get_file_content(file_path)
            
            # Create a temporary file if using S3
            if use_s3 and s3_client and not os.path.exists(file_path):
                temp_file = f"/tmp/{uuid.uuid4()}.pdf"
                with open(temp_file, 'wb') as f:
                    f.write(file_content)
                file_to_read = temp_file
            else:
                file_to_read = file_path
            
            # Read the PDF
            with open(file_to_read, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                # Check if the PDF is encrypted
                if pdf_reader.is_encrypted:
                    try:
                        pdf_reader.decrypt('')  # Try empty password
                    except:
                        text_parts.append("Error: PDF is encrypted and cannot be read")
                        return "\n".join(text_parts)
                
                # Extract text from each page
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text_parts.append(page.extract_text())
            
            # Clean up temporary file if created
            if 'temp_file' in locals() and os.path.exists(temp_file):
                os.remove(temp_file)
                
            return "\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {e}")
            return f"Error extracting text from PDF: {str(e)}"

    def _extract_text_from_word(self, file_path: str) -> str:
        """Extract text from a Word document.
        
        Args:
            file_path: Path to the Word document
            
        Returns:
            str: Extracted text
        """
        if not DOCX_AVAILABLE:
            return "Error: python-docx library not available for Word document extraction"
        
        try:
            # Get file content
            file_content = self.get_file_content(file_path)
            
            # Create a temporary file if using S3
            if use_s3 and s3_client and not os.path.exists(file_path):
                temp_file = f"/tmp/{uuid.uuid4()}.docx"
                with open(temp_file, 'wb') as f:
                    f.write(file_content)
                file_to_read = temp_file
            else:
                file_to_read = file_path
            
            # Read the Word document
            doc = docx.Document(file_to_read)
            text = "\n".join([para.text for para in doc.paragraphs])
            
            # Clean up temporary file if created
            if 'temp_file' in locals() and os.path.exists(temp_file):
                os.remove(temp_file)
                
            return text
            
        except Exception as e:
            logger.error(f"Error extracting text from Word document {file_path}: {e}")
            return f"Error extracting text from Word document: {str(e)}"

    def _extract_text_from_csv(self, file_path: str) -> str:
        """Extract text from a CSV file.
        
        Args:
            file_path: Path to the CSV file
            
        Returns:
            str: Extracted text
        """
        if not PANDAS_AVAILABLE:
            return "Error: pandas library not available for CSV extraction"
        
        try:
            # Get file content
            file_content = self.get_file_content(file_path)
            
            # Create a temporary file if using S3
            if use_s3 and s3_client and not os.path.exists(file_path):
                temp_file = f"/tmp/{uuid.uuid4()}.csv"
                with open(temp_file, 'wb') as f:
                    f.write(file_content)
                file_to_read = temp_file
            else:
                file_to_read = file_path
            
            # Read the CSV file
            df = pd.read_csv(file_to_read, encoding='utf-8', on_bad_lines='skip')
            
            # Convert to string representation
            text = df.to_string()
            
            # Clean up temporary file if created
            if 'temp_file' in locals() and os.path.exists(temp_file):
                os.remove(temp_file)
                
            return text
            
        except Exception as e:
            logger.error(f"Error extracting text from CSV {file_path}: {e}")
            return f"Error extracting text from CSV: {str(e)}"

    def _extract_text_from_excel(self, file_path: str) -> str:
        """Extract text from an Excel file.
        
        Args:
            file_path: Path to the Excel file
            
        Returns:
            str: Extracted text
        """
        if not PANDAS_AVAILABLE:
            return "Error: pandas library not available for Excel extraction"
        
        try:
            # Get file content
            file_content = self.get_file_content(file_path)
            
            # Create a temporary file if using S3
            if use_s3 and s3_client and not os.path.exists(file_path):
                temp_file = f"/tmp/{uuid.uuid4()}.xlsx"
                with open(temp_file, 'wb') as f:
                    f.write(file_content)
                file_to_read = temp_file
            else:
                file_to_read = file_path
            
            # Read the Excel file
            df = pd.read_excel(file_to_read, sheet_name=None)
            
            # Combine all sheets
            text_parts = []
            for sheet_name, sheet_df in df.items():
                text_parts.append(f"Sheet: {sheet_name}")
                text_parts.append(sheet_df.to_string())
                text_parts.append("")
            
            # Clean up temporary file if created
            if 'temp_file' in locals() and os.path.exists(temp_file):
                os.remove(temp_file)
                
            return "\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"Error extracting text from Excel {file_path}: {e}")
            return f"Error extracting text from Excel: {str(e)}"

    def _extract_pdf_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from a PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Dict[str, Any]: Extracted metadata
        """
        metadata = {}
        
        try:
            # Get file content
            file_content = self.get_file_content(file_path)
            
            # Create a temporary file if using S3
            if use_s3 and s3_client and not os.path.exists(file_path):
                temp_file = f"/tmp/{uuid.uuid4()}.pdf"
                with open(temp_file, 'wb') as f:
                    f.write(file_content)
                file_to_read = temp_file
            else:
                file_to_read = file_path
            
            # Read the PDF
            with open(file_to_read, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                metadata["page_count"] = len(pdf_reader.pages)
                
                # Extract document info
                if pdf_reader.metadata:
                    info = pdf_reader.metadata
                    # Convert to dict and handle Python types
                    metadata["title"] = info.get('/Title', '')
                    metadata["author"] = info.get('/Author', '')
                    metadata["subject"] = info.get('/Subject', '')
                    metadata["creator"] = info.get('/Creator', '')
                    metadata["producer"] = info.get('/Producer', '')
                
            # Clean up temporary file if created
            if 'temp_file' in locals() and os.path.exists(temp_file):
                os.remove(temp_file)
                
            return metadata
            
        except Exception as e:
            logger.error(f"Error extracting PDF metadata from {file_path}: {e}")
            return {"error": str(e)}

    def _extract_word_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from a Word document.
        
        Args:
            file_path: Path to the Word document
            
        Returns:
            Dict[str, Any]: Extracted metadata
        """
        metadata = {}
        
        try:
            # Get file content
            file_content = self.get_file_content(file_path)
            
            # Create a temporary file if using S3
            if use_s3 and s3_client and not os.path.exists(file_path):
                temp_file = f"/tmp/{uuid.uuid4()}.docx"
                with open(temp_file, 'wb') as f:
                    f.write(file_content)
                file_to_read = temp_file
            else:
                file_to_read = file_path
            
            # Read the Word document
            doc = docx.Document(file_to_read)
            
            # Extract metadata from core properties
            try:
                core_props = doc.core_properties
                metadata["title"] = core_props.title or ""
                metadata["author"] = core_props.author or ""
                metadata["comments"] = core_props.comments or ""
                metadata["keywords"] = core_props.keywords or ""
                metadata["subject"] = core_props.subject or ""
                metadata["created"] = core_props.created.isoformat() if core_props.created else ""
                metadata["modified"] = core_props.modified.isoformat() if core_props.modified else ""
            except Exception as e:
                logger.warning(f"Error extracting Word core properties: {e}")
            
            # Count paragraphs, tables, etc.
            metadata["paragraph_count"] = len(doc.paragraphs)
            metadata["table_count"] = len(doc.tables)
            metadata["section_count"] = len(doc.sections)
            
            # Clean up temporary file if created
            if 'temp_file' in locals() and os.path.exists(temp_file):
                os.remove(temp_file)
                
            return metadata
            
        except Exception as e:
            logger.error(f"Error extracting Word metadata from {file_path}: {e}")
            return {"error": str(e)}

    def _extract_csv_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from a CSV file.
        
        Args:
            file_path: Path to the CSV file
            
        Returns:
            Dict[str, Any]: Extracted metadata
        """
        metadata = {}
        
        try:
            # Get file content
            file_content = self.get_file_content(file_path)
            
            # Create a temporary file if using S3
            if use_s3 and s3_client and not os.path.exists(file_path):
                temp_file = f"/tmp/{uuid.uuid4()}.csv"
                with open(temp_file, 'wb') as f:
                    f.write(file_content)
                file_to_read = temp_file
            else:
                file_to_read = file_path
            
            # Read the CSV file
            df = pd.read_csv(file_to_read, encoding='utf-8', on_bad_lines='skip')
            
            # Extract basic statistics
            metadata["row_count"] = len(df)
            metadata["column_count"] = len(df.columns)
            metadata["columns"] = df.columns.tolist()
            
            # Clean up temporary file if created
            if 'temp_file' in locals() and os.path.exists(temp_file):
                os.remove(temp_file)
                
            return metadata
            
        except Exception as e:
            logger.error(f"Error extracting CSV metadata from {file_path}: {e}")
            return {"error": str(e)}

    def _extract_excel_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from an Excel file.
        
        Args:
            file_path: Path to the Excel file
            
        Returns:
            Dict[str, Any]: Extracted metadata
        """
        metadata = {}
        
        try:
            # Get file content
            file_content = self.get_file_content(file_path)
            
            # Create a temporary file if using S3
            if use_s3 and s3_client and not os.path.exists(file_path):
                temp_file = f"/tmp/{uuid.uuid4()}.xlsx"
                with open(temp_file, 'wb') as f:
                    f.write(file_content)
                file_to_read = temp_file
            else:
                file_to_read = file_path
            
            # Read the Excel file - get all sheets
            excel_file = pd.ExcelFile(file_to_read)
            
            # Get sheet names
            metadata["sheet_names"] = excel_file.sheet_names
            metadata["sheet_count"] = len(excel_file.sheet_names)
            
            # Get basic stats for each sheet
            sheet_stats = {}
            for sheet_name in excel_file.sheet_names:
                df = excel_file.parse(sheet_name)
                sheet_stats[sheet_name] = {
                    "row_count": len(df),
                    "column_count": len(df.columns),
                    "columns": df.columns.tolist()
                }
            
            metadata["sheets"] = sheet_stats
            
            # Clean up temporary file if created
            if 'temp_file' in locals() and os.path.exists(temp_file):
                os.remove(temp_file)
                
            return metadata
            
        except Exception as e:
            logger.error(f"Error extracting Excel metadata from {file_path}: {e}")
            return {"error": str(e)}

    def _analyze_pdf_structure(self, file_path: str) -> Dict[str, Any]:
        """Analyze the structure of a PDF document.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Dict[str, Any]: Document structure information
        """
        structure = {}
        
        try:
            # Get file content
            file_content = self.get_file_content(file_path)
            
            # Create a temporary file if using S3
            if use_s3 and s3_client and not os.path.exists(file_path):
                temp_file = f"/tmp/{uuid.uuid4()}.pdf"
                with open(temp_file, 'wb') as f:
                    f.write(file_content)
                file_to_read = temp_file
            else:
                file_to_read = file_path
            
            # Read the PDF
            with open(file_to_read, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                structure["total_pages"] = len(pdf_reader.pages)
                
                # Analyze first page to estimate document type
                if structure["total_pages"] > 0:
                    first_page_text = pdf_reader.pages[0].extract_text()
                    
                    # Simple heuristic checks for document type
                    if re.search(r'\b(contract|agreement)\b', first_page_text, re.IGNORECASE):
                        structure["estimated_type"] = "contract"
                    elif re.search(r'\b(invoice|bill|payment)\b', first_page_text, re.IGNORECASE):
                        structure["estimated_type"] = "invoice"
                    elif re.search(r'\b(report|analysis)\b', first_page_text, re.IGNORECASE):
                        structure["estimated_type"] = "report"
                    else:
                        structure["estimated_type"] = "general"
                
                # Check for common document parts
                structure["has_outline"] = hasattr(pdf_reader, 'outline') and pdf_reader.outline is not None
                
            # Clean up temporary file if created
            if 'temp_file' in locals() and os.path.exists(temp_file):
                os.remove(temp_file)
                
            return structure
            
        except Exception as e:
            logger.error(f"Error analyzing PDF structure for {file_path}: {e}")
            return {"error": str(e)}

    def _analyze_word_structure(self, file_path: str) -> Dict[str, Any]:
        """Analyze the structure of a Word document.
        
        Args:
            file_path: Path to the Word document
            
        Returns:
            Dict[str, Any]: Document structure information
        """
        structure = {}
        
        try:
            # Get file content
            file_content = self.get_file_content(file_path)
            
            # Create a temporary file if using S3
            if use_s3 and s3_client and not os.path.exists(file_path):
                temp_file = f"/tmp/{uuid.uuid4()}.docx"
                with open(temp_file, 'wb') as f:
                    f.write(file_content)
                file_to_read = temp_file
            else:
                file_to_read = file_path
            
            # Read the Word document
            doc = docx.Document(file_to_read)
            
            # Analyze document structure
            structure["paragraph_count"] = len(doc.paragraphs)
            structure["table_count"] = len(doc.tables)
            structure["section_count"] = len(doc.sections)
            
            # Count headings by level
            heading_counts = {}
            for paragraph in doc.paragraphs:
                if paragraph.style.name.startswith('Heading'):
                    heading_level = paragraph.style.name.replace('Heading ', '')
                    if heading_level.isdigit():
                        heading_counts[f"heading_{heading_level}"] = heading_counts.get(f"heading_{heading_level}", 0) + 1
            
            structure["headings"] = heading_counts
            
            # Estimate document type based on content
            all_text = "\n".join([p.text for p in doc.paragraphs])
            if re.search(r'\b(contract|agreement)\b', all_text, re.IGNORECASE):
                structure["estimated_type"] = "contract"
            elif re.search(r'\b(invoice|bill|payment)\b', all_text, re.IGNORECASE):
                structure["estimated_type"] = "invoice"
            elif re.search(r'\b(report|analysis)\b', all_text, re.IGNORECASE):
                structure["estimated_type"] = "report"
            else:
                structure["estimated_type"] = "general"
            
            # Clean up temporary file if created
            if 'temp_file' in locals() and os.path.exists(temp_file):
                os.remove(temp_file)
                
            return structure
            
        except Exception as e:
            logger.error(f"Error analyzing Word structure for {file_path}: {e}")
            return {"error": str(e)}

    def _analyze_csv_structure(self, file_path: str) -> Dict[str, Any]:
        """Analyze the structure of a CSV file.
        
        Args:
            file_path: Path to the CSV file
            
        Returns:
            Dict[str, Any]: Document structure information
        """
        structure = {}
        
        try:
            # Get file content
            file_content = self.get_file_content(file_path)
            
            # Create a temporary file if using S3
            if use_s3 and s3_client and not os.path.exists(file_path):
                temp_file = f"/tmp/{uuid.uuid4()}.csv"
                with open(temp_file, 'wb') as f:
                    f.write(file_content)
                file_to_read = temp_file
            else:
                file_to_read = file_path
            
            # Read the CSV file
            df = pd.read_csv(file_to_read, encoding='utf-8', on_bad_lines='skip')
            
            # Analyze structure
            structure["row_count"] = len(df)
            structure["column_count"] = len(df.columns)
            structure["columns"] = df.columns.tolist()
            
            # Get column types
            column_types = {}
            for column in df.columns:
                if pd.api.types.is_numeric_dtype(df[column]):
                    if pd.api.types.is_integer_dtype(df[column]):
                        column_types[column] = "integer"
                    else:
                        column_types[column] = "float"
                elif pd.api.types.is_datetime64_dtype(df[column]):
                    column_types[column] = "datetime"
                else:
                    column_types[column] = "string"
            
            structure["column_types"] = column_types
            
            # Guess if the CSV is a table or just data
            # Heuristic: If first column has sequential values or indices, it's likely a table
            is_sequential = False
            if structure["row_count"] > 1 and structure["column_count"] > 1:
                first_col = df.iloc[:, 0]
                if pd.api.types.is_numeric_dtype(first_col):
                    # Check if values are sequential or increasing
                    diffs = first_col.diff().dropna()
                    if all(diffs > 0) and (diffs.nunique() <= 2):  # Allow for some small variation
                        is_sequential = True
            
            structure["likely_table"] = is_sequential
            
            # Clean up temporary file if created
            if 'temp_file' in locals() and os.path.exists(temp_file):
                os.remove(temp_file)
                
            return structure
            
        except Exception as e:
            logger.error(f"Error analyzing CSV structure for {file_path}: {e}")
            return {"error": str(e)}

    def _analyze_excel_structure(self, file_path: str) -> Dict[str, Any]:
        """Analyze the structure of an Excel file.
        
        Args:
            file_path: Path to the Excel file
            
        Returns:
            Dict[str, Any]: Document structure information
        """
        structure = {}
        
        try:
            # Get file content
            file_content = self.get_file_content(file_path)
            
            # Create a temporary file if using S3
            if use_s3 and s3_client and not os.path.exists(file_path):
                temp_file = f"/tmp/{uuid.uuid4()}.xlsx"
                with open(temp_file, 'wb') as f:
                    f.write(file_content)
                file_to_read = temp_file
            else:
                file_to_read = file_path
            
            # Read the Excel file - get all sheets
            excel_file = pd.ExcelFile(file_to_read)
            
            # Get sheet names
            structure["sheet_names"] = excel_file.sheet_names
            structure["sheet_count"] = len(excel_file.sheet_names)
            
            # Analyze each sheet
            sheet_structures = {}
            for sheet_name in excel_file.sheet_names:
                df = excel_file.parse(sheet_name)
                
                sheet_info = {
                    "row_count": len(df),
                    "column_count": len(df.columns),
                    "columns": df.columns.tolist()
                }
                
                # Get column types
                column_types = {}
                for column in df.columns:
                    if pd.api.types.is_numeric_dtype(df[column]):
                        if pd.api.types.is_integer_dtype(df[column]):
                            column_types[str(column)] = "integer"
                        else:
                            column_types[str(column)] = "float"
                    elif pd.api.types.is_datetime64_dtype(df[column]):
                        column_types[str(column)] = "datetime"
                    else:
                        column_types[str(column)] = "string"
                
                sheet_info["column_types"] = column_types
                
                # Check if the sheet appears to be a data table or something else
                # Heuristic: If it has a header row and consistent data types in columns
                has_header = True  # Assume pandas correctly identified the header
                is_data_table = has_header and (len(df) > 0) and (len(df.columns) > 0)
                sheet_info["likely_data_table"] = is_data_table
                
                sheet_structures[sheet_name] = sheet_info
            
            structure["sheets"] = sheet_structures
            
            # Clean up temporary file if created
            if 'temp_file' in locals() and os.path.exists(temp_file):
                os.remove(temp_file)
                
            return structure
            
        except Exception as e:
            logger.error(f"Error analyzing Excel structure for {file_path}: {e}")
            return {"error": str(e)}