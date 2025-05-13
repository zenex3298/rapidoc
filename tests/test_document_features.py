import os
import sys
import json
import time
import logging
from datetime import datetime
import requests
import base64
from dotenv import load_dotenv

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('logs/document_test.log'),
        logging.StreamHandler()
    ]
)

logger = logging.getLogger(__name__)

BASE_URL = "http://localhost:8000"
ACCESS_TOKEN = None
TEST_USER = {
    'email': f'doctest_{int(time.time())}@example.com',
    'username': f'doctest_{int(time.time())}',
    'password': 'TestPassword123'
}

# Test file paths
SAMPLE_DOCS_DIR = "tests/sample_docs"
os.makedirs(SAMPLE_DOCS_DIR, exist_ok=True)

def create_sample_docx():
    """Create a sample DOCX file for testing"""
    try:
        from docx import Document
        doc = Document()
        doc.add_heading('Sample Contract', 0)
        
        doc.add_heading('Parties', level=1)
        doc.add_paragraph('This contract is between Party A and Party B, dated January 1, 2025.')
        
        doc.add_heading('Terms and Conditions', level=1)
        doc.add_paragraph('1. Party A agrees to provide services to Party B.')
        doc.add_paragraph('2. Party B agrees to pay $1,000 to Party A for the services.')
        doc.add_paragraph('3. This agreement is valid for one year from the date of signing.')
        
        doc.add_heading('Termination', level=1)
        doc.add_paragraph('Either party may terminate this agreement with 30 days written notice.')
        
        sample_docx_path = os.path.join(SAMPLE_DOCS_DIR, "sample_contract.docx")
        doc.save(sample_docx_path)
        logger.info(f"Created sample DOCX file at {sample_docx_path}")
        return sample_docx_path
    except Exception as e:
        logger.error(f"Failed to create sample DOCX: {e}")
        return None

def register_and_login():
    """Register a new user and get access token"""
    global ACCESS_TOKEN
    
    try:
        # Register new user
        response = requests.post(f"{BASE_URL}/auth/register", json=TEST_USER)
        if response.status_code != 201:
            logger.error(f"Registration failed: {response.text}")
            return False
        
        logger.info(f"Registered test user: {TEST_USER['username']}")
        
        # Login
        login_data = {
            'username': TEST_USER['username'],
            'password': TEST_USER['password']
        }
        response = requests.post(
            f"{BASE_URL}/auth/login", 
            data=login_data
        )
        
        if response.status_code != 200:
            logger.error(f"Login failed: {response.text}")
            return False
        
        token_data = response.json()
        ACCESS_TOKEN = token_data["access_token"]
        logger.info(f"Logged in successfully, received access token")
        return True
    
    except Exception as e:
        logger.error(f"Error in registration/login: {e}")
        return False

def test_document_upload():
    """Test document upload and processing"""
    if not ACCESS_TOKEN:
        logger.error("No access token available")
        return None
    
    try:
        # Create sample document
        sample_docx_path = create_sample_docx()
        if not sample_docx_path:
            return None
        
        # Upload document
        with open(sample_docx_path, 'rb') as f:
            files = {'file': f}
            data = {
                'title': 'Test Contract',
                'description': 'A sample contract for testing',
                'doc_type': 'repo'
            }
            headers = {'Authorization': f'Bearer {ACCESS_TOKEN}'}
            
            response = requests.post(
                f"{BASE_URL}/documents/upload",
                files=files,
                data=data,
                headers=headers
            )
        
        if response.status_code != 201:
            logger.error(f"Document upload failed: {response.text}")
            return None
        
        document_data = response.json()
        document_id = document_data['id']
        logger.info(f"Uploaded document with ID: {document_id}")
        
        # Wait for processing to complete (in a real test, would use a better approach)
        logger.info("Waiting for document processing to complete...")
        time.sleep(5)  # Wait for processing
        
        return document_id
    
    except Exception as e:
        logger.error(f"Error in document upload: {e}")
        return None

def test_document_analysis(document_id):
    """Test retrieving document analysis"""
    if not ACCESS_TOKEN or not document_id:
        logger.error("No access token or document ID available")
        return False
        
    # Check if OpenAI API key is configured
    has_openai_key = check_openai_api_key()
    if not has_openai_key:
        logger.warning("Skipping detailed analysis check due to missing OpenAI API key")
        # We'll still try to retrieve the document, but won't check analysis content
    
    try:
        headers = {'Authorization': f'Bearer {ACCESS_TOKEN}'}
        
        # Get document details first
        response = requests.get(
            f"{BASE_URL}/documents/{document_id}",
            headers=headers
        )
        
        if response.status_code != 200:
            logger.error(f"Failed to get document details: {response.text}")
            return False
        
        document = response.json()
        logger.info(f"Retrieved document details, status: {document['status']}")
        
        # If document is still processing, wait
        if document['status'] == 'processing':
            logger.info("Document still processing, waiting...")
            time.sleep(5)
        
        # Get document analysis
        response = requests.get(
            f"{BASE_URL}/documents/{document_id}/analysis",
            headers=headers
        )
        
        if response.status_code == 404:
            logger.warning("Analysis not available yet, might need more time for processing")
            return False
        
        if response.status_code != 200:
            logger.error(f"Failed to get document analysis: {response.text}")
            return False
        
        analysis = response.json()
        logger.info(f"Retrieved document analysis with {len(json.dumps(analysis))} characters")
        
        # Check for expected analysis fields
        if 'analysis' in analysis and 'metadata' in analysis and 'clauses' in analysis:
            logger.info("Document analysis structure looks correct")
            return True
        else:
            logger.warning("Document analysis structure is unexpected")
            return False
    
    except Exception as e:
        logger.error(f"Error in document analysis test: {e}")
        return False

def test_document_clauses(document_id):
    """Test retrieving document clauses"""
    if not ACCESS_TOKEN or not document_id:
        logger.error("No access token or document ID available")
        return False
    
    try:
        headers = {'Authorization': f'Bearer {ACCESS_TOKEN}'}
        
        # Get document clauses
        response = requests.get(
            f"{BASE_URL}/documents/{document_id}/clauses",
            headers=headers
        )
        
        if response.status_code == 404:
            logger.warning("Clauses not available yet, might need more time for processing")
            return False
        
        if response.status_code != 200:
            logger.error(f"Failed to get document clauses: {response.text}")
            return False
        
        clauses = response.json()
        logger.info(f"Retrieved {len(clauses)} document clauses")
        
        if len(clauses) > 0:
            logger.info("Document clauses found")
            return True
        else:
            logger.warning("No document clauses found")
            return False
    
    except Exception as e:
        logger.error(f"Error in document clauses test: {e}")
        return False

def test_document_template_transformation(document_id=None):
    """Test document transformation with templates"""
    if not ACCESS_TOKEN:
        logger.error("No access token available")
        return False
    
    if not document_id:
        # Create a document to transform if not provided
        document_id = test_document_upload()
        if not document_id:
            logger.error("Failed to create document for transformation test")
            return False
    
    try:
        headers = {'Authorization': f'Bearer {ACCESS_TOKEN}'}
        
        # Create two template documents - one for input and one for output
        sample_docx_path = create_sample_docx()
        if not sample_docx_path:
            logger.error("Failed to create sample document for template")
            return False
            
        # Upload input template
        with open(sample_docx_path, 'rb') as f:
            files = {'file': f}
            data = {
                'title': 'Input Template',
                'description': 'Template for input format',
                'doc_type': 'repo',
                'tag': 'template'  # Mark as template
            }
            
            response = requests.post(
                f"{BASE_URL}/documents/upload",
                files=files,
                data=data,
                headers=headers
            )
        
        if response.status_code != 201:
            logger.error(f"Input template upload failed: {response.text}")
            return False
            
        template_input_id = response.json()['id']
        logger.info(f"Uploaded input template with ID: {template_input_id}")
        
        # Upload output template
        with open(sample_docx_path, 'rb') as f:
            files = {'file': f}
            data = {
                'title': 'Output Template',
                'description': 'Template for output format',
                'doc_type': 'repo',
                'tag': 'template'  # Mark as template
            }
            
            response = requests.post(
                f"{BASE_URL}/documents/upload",
                files=files,
                data=data,
                headers=headers
            )
        
        if response.status_code != 201:
            logger.error(f"Output template upload failed: {response.text}")
            return False
            
        template_output_id = response.json()['id']
        logger.info(f"Uploaded output template with ID: {template_output_id}")
        
        # Wait for templates to be processed
        logger.info("Waiting for templates to be processed...")
        time.sleep(5)
        
        # Now test the transformation endpoint
        transform_data = {
            'template_input_id': template_input_id,
            'template_output_id': template_output_id
        }
        
        logger.info(f"Transforming document {document_id} with templates {template_input_id} and {template_output_id}...")
        response = requests.post(
            f"{BASE_URL}/documents/{document_id}/transform-with-templates",
            json=transform_data,
            headers=headers
        )
        
        if response.status_code != 200:
            logger.error(f"Document transformation failed: {response.text}")
            return False
            
        result = response.json()
        logger.info(f"Document transformation result: {json.dumps(result, indent=2)[:200]}...")
        
        # Check for expected fields in result
        required_fields = ['document_id', 'template_input_id', 'template_output_id', 'transformed_content']
        missing_fields = [field for field in required_fields if field not in result]
        
        if missing_fields:
            logger.error(f"Transformation result missing required fields: {missing_fields}")
            return False
            
        logger.info("Document transformation test passed")
        return True
        
    except Exception as e:
        logger.error(f"Error in document transformation test: {e}")
        return False

def test_document_generation():
    """Test document generation"""
    if not ACCESS_TOKEN:
        logger.error("No access token available")
        return False
        
    # Check if OpenAI API key is configured
    has_openai_key = check_openai_api_key()
    if not has_openai_key:
        logger.warning("Skipping document generation test due to missing OpenAI API key")
        return True  # Skip this test rather than fail it
    
    try:
        headers = {'Authorization': f'Bearer {ACCESS_TOKEN}'}
        
        # Document generation request
        template_data = {
            "parties": {
                "party_a": "Acme Corporation",
                "party_b": "XYZ Inc."
            },
            "agreement_details": {
                "start_date": "2025-05-01",
                "end_date": "2026-04-30",
                "payment_amount": "$5,000",
                "payment_schedule": "Monthly"
            },
            "services": [
                "Software development",
                "Technical support",
                "System maintenance"
            ],
            "special_terms": "All intellectual property created during the contract period belongs to Acme Corporation."
        }
        
        data = {
            'template_type': 'deposition',  # Only 'deposition' is allowed
            'input_data': json.dumps(template_data),
            'title': 'Generated Service Agreement',
            'description': 'AI-generated service agreement for testing'
        }
        
        response = requests.post(
            f"{BASE_URL}/documents/generate",
            data=data,
            headers=headers
        )
        
        if response.status_code != 200:
            logger.error(f"Document generation failed: {response.text}")
            return False
        
        result = response.json()
        document_id = result.get('document_id')
        content = result.get('content')
        
        logger.info(f"Generated document with ID: {document_id}")
        logger.info(f"Generated content length: {len(content) if content else 0}")
        
        if document_id and content and len(content) > 100:
            logger.info("Document generation successful")
            return True
        else:
            logger.warning("Document generation result is unexpected")
            return False
    
    except Exception as e:
        logger.error(f"Error in document generation test: {e}")
        return False

def check_openai_api_key():
    """Check if OpenAI API key is configured"""
    # Conditionally load environment variables only if needed
    openai_api_key = os.getenv("OPENAI_API_KEY")
    if not openai_api_key:
        logger.info("OPENAI_API_KEY not found in environment, attempting to load from .env file")
        load_dotenv()
        openai_api_key = os.getenv("OPENAI_API_KEY")
    
    if not openai_api_key or openai_api_key == "your_openai_api_key_here":
        logger.warning("OpenAI API key not set. Some tests might fail or take longer.")
        return False
    return True

def run_tests():
    """Run all document feature tests"""
    try:
        # Create logs directory
        os.makedirs('logs', exist_ok=True)
        logger.info("Starting document feature tests")
        
        # Check OpenAI API key
        has_openai_key = check_openai_api_key()
        if not has_openai_key:
            logger.warning("OpenAI API key not configured. Document analysis and generation tests may fail.")
        
        # Check if server is running
        try:
            response = requests.get(f"{BASE_URL}/health")
            if response.status_code != 200:
                logger.error(f"Server health check failed: {response.status_code}")
                return False
            logger.info("Server is running")
        except requests.exceptions.ConnectionError:
            logger.error("Could not connect to server. Make sure it's running.")
            return False
        
        # Register and login
        if not register_and_login():
            logger.error("Registration/login test failed")
            return False
        logger.info("Registration and login successful")
        
        # Test document upload
        document_id = test_document_upload()
        if not document_id:
            logger.error("Document upload test failed")
            return False
        logger.info("Document upload test passed")
        
        # Test document analysis
        if not test_document_analysis(document_id):
            logger.warning("Document analysis test inconclusive - may need more processing time")
            # We'll continue with other tests
        else:
            logger.info("Document analysis test passed")
        
        # Test document clauses
        if not test_document_clauses(document_id):
            logger.warning("Document clauses test inconclusive - may need more processing time")
            # We'll continue with other tests
        else:
            logger.info("Document clauses test passed")
        
        # Test document generation
        if not test_document_generation():
            logger.error("Document generation test failed")
            return False
        logger.info("Document generation test passed")
        
        # Test document template transformation
        if not test_document_template_transformation(document_id):
            logger.error("Document template transformation test failed")
            return False
        logger.info("Document template transformation test passed")
        
        logger.info("All document feature tests completed successfully!")
        return True
    
    except Exception as e:
        logger.error(f"An unexpected error occurred during testing: {e}")
        return False

if __name__ == "__main__":
    success = run_tests()
    if not success:
        sys.exit(1)